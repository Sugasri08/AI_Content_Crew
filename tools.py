import io
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document

def create_pdf(text):
    """Generates a raw PDF byte stream from standard text output."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    story = []
    
    styles = getSampleStyleSheet()
    # Create custom copy style supporting text wraps
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        spaceAfter=10
    )
    
    # Process text fragments separated by breaks safely
    lines = text.split("\n")
    for line in lines:
        if line.strip():
            clean_line = line.replace("<", "&lt;").replace(">", "&gt;") # Escape HTML bindings
            story.append(Paragraph(clean_line, body_style))
        else:
            story.append(Spacer(1, 10))
            
    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def create_docx(text):
    """Generates a raw DOCX byte stream from standard text output."""
    doc = Document()
    lines = text.split("\n")
    for line in lines:
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()